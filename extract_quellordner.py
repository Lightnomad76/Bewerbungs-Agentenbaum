# -*- coding: utf-8 -*-
"""QuellordnerExtractor — deterministischer, OFFLINE, READ-ONLY Text-Extraktor fuer
den persoenlichen Bewerbungs-Quellordner (jahrelange Zeugnisse / Lebenslaeufe /
Anschreiben / Eigenbemuehungs-Nachweise).

Pipeline-Rolle: liefert die *echten* biografischen Fakten (Firmen, Zeitraeume,
Zeugnis-Inhalte) als zusaetzliche Wahrheitsquelle fuer den FactGroundingAgent
(`factground.sammle_fakten`). Bisher waren die Stammdaten von Hand reingefuettert;
dieser Extraktor zieht sie maschinell aus den Originaldokumenten.

HARTE INVARIANTEN
-----------------
1. READ-ONLY am Quellordner: liest ausschliesslich; schreibt NIE in den Quellordner
   (verschiebt/aendert/loescht nichts). Der einzige Schreibvorgang ist der lokale
   Corpus-Output, und der DARF NICHT im Quellordner liegen (wird hart geprueft).
2. OFFLINE / KEINE API / kostenlos: nur lokale Libs (python-docx, pdfplumber) und das
   lokal vorhandene `antiword` (Git-mingw64) fuer alt-`.doc`. Kein Netz, keine LLM.
3. PII: Quellinhalte sind personenbezogen -> der erzeugte Corpus ist lokal/gitignored,
   NIE committen/pushen (public repo).

Unterstuetzte Formate (Coverage-Probe 2026-06-28: 58/58 textfuehrende Dateien sauber):
  .odt  -> ZIP/content.xml (reines Python, keine Dependency)
  .docx -> python-docx
  .pdf  -> pdfplumber (Text-Layer; reine Scans liefern leer -> Status SKIP)
  .doc  -> antiword (Subprozess; Output cp1252). Fehlt antiword -> Status SKIP.
Bild/Video/PSD (.jpg/.png/.psd/.mp4) werden bewusst ignoriert (kein OCR = keine neue
Dependency); ihr Inhalt liegt ohnehin zusaetzlich als Text-PDF vor.

CLI:  py -3.11 extract_quellordner.py <quellordner> [--out corpus.txt] [--json]
API:  from extract_quellordner import extrahiere, corpus_text, schreibe_corpus
      texte = extrahiere(quellordner)        # [(relpath, text, status), ...]
      corpus = corpus_text(quellordner)      # ein konsolidierter str
      stats  = schreibe_corpus(quellordner, "profile/quellordner_corpus.txt")
"""
from __future__ import annotations

import os
import re
import sys
import json
import zipfile
import subprocess

# Reihenfolge irrelevant; nur diese Suffixe werden ueberhaupt angefasst.
TEXT_SUFFIXE = (".odt", ".docx", ".pdf", ".doc")

# Unterhalb dieser Wortzahl gilt eine Datei als Scan/leer -> Status SKIP statt OK.
MIN_WORTE = 20

# antiword-Output ist cp1252 (empirisch verifiziert 2026-06-28, 0 Ersatzzeichen).
_ANTIWORD_ENC = "cp1252"


# ---------------------------------------------------------------------------
# Einzel-Extraktoren (jeder gibt reinen Text zurueck oder wirft)
# ---------------------------------------------------------------------------


def _odt_text(pfad: str) -> str:
    with zipfile.ZipFile(pfad) as z:
        xml = z.read("content.xml").decode("utf-8", "replace")
    # Absatz-/Zeilen-Tags zu Umbruch, restliche Tags zu Leerzeichen.
    xml = re.sub(r"</text:(?:p|h)>", "\n", xml)
    text = re.sub(r"<[^>]+>", " ", xml)
    # XML-Entities minimal aufloesen.
    for a, b in (("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"),
                 ("&apos;", "'"), ("&quot;", '"')):
        text = text.replace(a, b)
    return re.sub(r"[ \t]+", " ", text)


def _docx_text(pfad: str) -> str:
    import docx  # lokal, damit der Import-Fehler nur .docx betrifft
    d = docx.Document(pfad)
    teile = [p.text for p in d.paragraphs]
    # Auch Tabellen-Zellen (Lebenslaeufe nutzen oft 2-Spalten-Tabellen).
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                teile.append(cell.text)
    return "\n".join(t for t in teile if t)


def _pdf_text(pfad: str) -> str:
    import pdfplumber
    with pdfplumber.open(pfad) as pdf:
        return "\n".join((pg.extract_text() or "") for pg in pdf.pages)


def _doc_text(pfad: str) -> str:
    # antiword schreibt nach stdout; wir lesen NUR (kein Schreibzugriff auf den Quellordner).
    try:
        out = subprocess.run(
            ["antiword", "-w", "0", pfad],
            capture_output=True, timeout=60,
        )
    except FileNotFoundError:
        raise RuntimeError("antiword nicht gefunden (PATH) — .doc uebersprungen")
    if out.returncode != 0:
        raise RuntimeError("antiword rc=%d: %s"
                           % (out.returncode, out.stderr.decode(_ANTIWORD_ENC, "replace")[:120]))
    return out.stdout.decode(_ANTIWORD_ENC, "replace")


_EXTRAKTOR = {
    ".odt": _odt_text,
    ".docx": _docx_text,
    ".pdf": _pdf_text,
    ".doc": _doc_text,
}


# ---------------------------------------------------------------------------
# Walk + Konsolidierung
# ---------------------------------------------------------------------------


def _iter_dateien(basis: str):
    """Alle textfuehrenden Dateien unter basis (rekursiv), sortiert + deterministisch."""
    treffer = []
    for wurzel, _dirs, dateien in os.walk(basis):
        for name in dateien:
            if os.path.splitext(name)[1].lower() in TEXT_SUFFIXE:
                treffer.append(os.path.join(wurzel, name))
    return sorted(treffer)


def extrahiere(basis: str):
    """Quellordner read-only auslesen.

    Returns: Liste von (relpath, text, status) mit status in {"OK","SKIP","ERR"}.
    "SKIP" = Datei lesbar, aber < MIN_WORTE (Scan/leer) ODER Extraktor nicht verfuegbar.
    "ERR"  = Extraktion fehlgeschlagen (Datei korrupt o. ae.).
    """
    if not os.path.isdir(basis):
        raise NotADirectoryError("Quellordner nicht gefunden: " + basis)
    ergebnis = []
    for pfad in _iter_dateien(basis):
        rel = os.path.relpath(pfad, basis)
        suf = os.path.splitext(pfad)[1].lower()
        try:
            text = _EXTRAKTOR[suf](pfad)
        except RuntimeError as e:  # antiword fehlt / nonzero -> SKIP, kein Abbruch
            ergebnis.append((rel, "", "SKIP:" + str(e)))
            continue
        except Exception as e:  # noqa: BLE001 — korrupte Datei darf den Lauf nicht killen
            ergebnis.append((rel, "", "ERR:" + str(e)))
            continue
        if len(text.split()) < MIN_WORTE:
            ergebnis.append((rel, text, "SKIP:zu kurz (<%d Worte, Scan/leer?)" % MIN_WORTE))
        else:
            ergebnis.append((rel, text, "OK"))
    return ergebnis


def corpus_text(basis: str) -> str:
    """Konsolidierter Text aller OK-Dateien, je Datei mit Pfad-Header (zur Nachverfolgung)."""
    bloecke = []
    for rel, text, status in extrahiere(basis):
        if status == "OK":
            bloecke.append("### " + rel + "\n" + text.strip())
    return "\n\n".join(bloecke)


def _guard_out_pfad(basis: str, out_pfad: str) -> None:
    """READ-ONLY-Schutz: der Output darf NICHT im Quellordner liegen."""
    b = os.path.abspath(basis)
    o = os.path.abspath(out_pfad)
    if o == b or o.startswith(b + os.sep):
        raise ValueError("READ-ONLY-Verletzung: Corpus-Output liegt im Quellordner (" + o + ")")


def schreibe_corpus(basis: str, out_pfad: str) -> dict:
    """Corpus in out_pfad schreiben (out MUSS ausserhalb des Quellordners liegen).
    Gibt Statistik zurueck. Quellordner wird NUR gelesen."""
    _guard_out_pfad(basis, out_pfad)
    eintraege = extrahiere(basis)
    bloecke = ["### " + rel + "\n" + text.strip()
               for rel, text, status in eintraege if status == "OK"]
    corpus = "\n\n".join(bloecke)
    out_dir = os.path.dirname(os.path.abspath(out_pfad))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    with open(out_pfad, "w", encoding="utf-8") as fh:
        fh.write(corpus)
    return _stats(eintraege, len(corpus))


def _stats(eintraege, corpus_len: int) -> dict:
    ok = sum(1 for _, _, s in eintraege if s == "OK")
    skip = sum(1 for _, _, s in eintraege if s.startswith("SKIP"))
    err = sum(1 for _, _, s in eintraege if s.startswith("ERR"))
    return {"dateien": len(eintraege), "ok": ok, "skip": skip, "err": err,
            "corpus_zeichen": corpus_len}


# ---------------------------------------------------------------------------
# Report / CLI
# ---------------------------------------------------------------------------


def report(basis: str, eintraege) -> str:
    st = _stats(eintraege, sum(len(t) for _, t, s in eintraege if s == "OK"))
    out = ["=== QuellordnerExtractor (READ-ONLY) ===", "Quelle: " + basis,
           "Dateien: %d  (OK %d / SKIP %d / ERR %d)"
           % (st["dateien"], st["ok"], st["skip"], st["err"]), ""]
    for rel, text, status in eintraege:
        kurz = status if status == "OK" else status.split(":", 1)[0]
        worte = len(text.split())
        out.append("  [%-4s] %5d W  %s" % (kurz, worte, rel))
        if status.startswith(("ERR", "SKIP")) and ":" in status:
            out.append("           -> " + status.split(":", 1)[1])
    return "\n".join(out)


def main(argv):
    args = [a for a in argv[1:] if not a.startswith("--")]
    flags = {a for a in argv[1:] if a.startswith("--")}
    out_pfad = None
    for a in argv[1:]:
        if a.startswith("--out="):
            out_pfad = a.split("=", 1)[1]
    if not args:
        print("Aufruf: py -3.11 extract_quellordner.py <quellordner> [--out=corpus.txt] [--json]",
              file=sys.stderr)
        return 2
    basis = args[0]
    try:
        eintraege = extrahiere(basis)
    except (NotADirectoryError, OSError) as e:
        print("FEHLER: " + str(e), file=sys.stderr)
        return 2

    if out_pfad:
        try:
            st = schreibe_corpus(basis, out_pfad)
        except ValueError as e:
            print("FEHLER: " + str(e), file=sys.stderr)
            return 2
        print("Corpus geschrieben -> " + out_pfad
              + "  (%d Zeichen, %d/%d Dateien OK)" % (st["corpus_zeichen"], st["ok"], st["dateien"]))

    if "--json" in flags:
        st = _stats(eintraege, sum(len(t) for _, t, s in eintraege if s == "OK"))
        print(json.dumps({"basis": basis, "stats": st,
                          "dateien": [{"datei": r, "worte": len(t.split()), "status": s}
                                      for r, t, s in eintraege]},
                         ensure_ascii=False, indent=2))
    else:
        print(report(basis, eintraege))
    # exit 1, wenn gar nichts Brauchbares extrahiert wurde (Konfig-/Pfad-Fehler-Signal)
    return 0 if any(s == "OK" for _, _, s in eintraege) else 1


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    sys.exit(main(sys.argv))
