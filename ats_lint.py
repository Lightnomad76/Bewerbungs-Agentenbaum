# -*- coding: utf-8 -*-
"""ATS-Linter — deterministische, offline Pruefung eines CV-.docx auf
ATS-Parsing-Risiken (Applicant Tracking System).

Pipeline-Rolle (Agenten-Roadmap E-C, Teil 3 / Abschnitt C): prueft das vom Generator
erzeugte Lebenslauf-.docx mechanisch auf bekannte ATS-Stolpersteine, BEVOR es in ein
Konzern-ATS (z. B. SAP SuccessFactors) hochgeladen wird. KEINE LLM-API, rein
strukturell ueber python-docx + OOXML.

Grundsatz wie critic/factground/jdparser: der Linter FLAGGT, aendert nichts
(read-only Scope). Er ersetzt KEINEN echten Parser-Durchlauf (§3.6a bleibt offen) —
er faengt nur die in der Recherche belegten, mechanisch erkennbaren Risiken vorab ab.

Belegte Risiken (state/agenten_roadmap.md Abschnitt C):
- FEHLER: Tabellen (Taleo scrambelt Zellen / Workday merged), mehrspaltiges
          Seitenlayout, Textboxen, Kontaktdaten in Kopf-/Fusszeile (ATS ignoriert),
          praktisch kein extrahierbarer Text (Scan-/Bild-CV).
- WARNUNG: eingebettete Bilder/Grafiken (ATS ueberspringt sie), keine erkennbaren
           Standard-Abschnittsueberschriften, keine Kontaktdaten als Plain-Text im
           Body.

EHRLICHE GRENZE: "randlos" schuetzt NICHT — der Fix ist Single-Column, nicht eine
unsichtbare Tabelle. Welcher Anteil der Ziel-Stellen real durch ein ATS laeuft, ist
unbekannt (untested); der Linter sagt nur, ob das Dokument ATS-robust *aufgebaut* ist.

CLI:   py -3.11 ats_lint.py <cv.docx> [--json]   (exit 1 bei FEHLER / 0 sonst)
API:   from ats_lint import pruefe, pruefe_datei, hat_fehler
       res = pruefe_datei("Lebenslauf_ATS.docx"); hat_fehler(res) -> bool
"""
from __future__ import annotations

import re
import sys
import json

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:  # pragma: no cover
    Document = None
    qn = None

# ---------------------------------------------------------------------------
# Muster + Vokabular
# ---------------------------------------------------------------------------

EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
TEL_RE = re.compile(r"(tel\.?|telefon|mobil|handy)\b", re.IGNORECASE)
TEL_NUM_RE = re.compile(r"\b0\d[\d\s/().+-]{6,}\d\b")

# Standard-CV-Abschnittsueberschriften (lowercase, Teilstring-Match).
STANDARD_HEADINGS = [
    "berufserfahrung", "berufliche erfahrung", "beruflicher werdegang", "werdegang",
    "ausbildung", "schulausbildung", "schulbildung", "studium",
    "kenntnisse", "fähigkeiten", "faehigkeiten", "fachkenntnisse", "qualifikation",
    "edv", "it-kenntnisse", "weiterbildung", "fortbildung",
    "persönliche daten", "persoenliche daten", "sprachen", "praktika", "interessen",
]

MIN_TEXT_LEN = 200  # weniger extrahierbarer Text -> Verdacht Scan/Bild-CV


# ---------------------------------------------------------------------------
# Hilfsfunktionen
# ---------------------------------------------------------------------------


def _finding(kategorie, schwere, nachricht, fundstelle=None):
    return {"kategorie": kategorie, "schwere": schwere, "nachricht": nachricht,
            "fundstelle": fundstelle}


def _alle_absaetze(doc):
    """Alle Absatz-Texte: Body + Tabellen-Zellen (Headings koennen in Zellen liegen)."""
    texte = [p.text for p in doc.paragraphs]
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                texte.extend(p.text for p in cell.paragraphs)
    return [t for t in texte if t and t.strip()]


def _kopf_fuss_text(doc) -> str:
    teile = []
    for sec in doc.sections:
        for behaelter in (sec.header, sec.footer,
                          sec.first_page_header, sec.first_page_footer,
                          sec.even_page_header, sec.even_page_footer):
            try:
                teile.extend(p.text for p in behaelter.paragraphs)
            except Exception:
                pass
    return "\n".join(t for t in teile if t)


def _spalten_anzahl(doc) -> int:
    """Maximale Spaltenzahl ueber alle Sektionen (mehrspaltiges Layout = Risiko)."""
    maxn = 1
    for sec in doc.sections:
        cols = sec._sectPr.find(qn("w:cols"))
        if cols is not None:
            num = cols.get(qn("w:num"))
            if num and num.isdigit():
                maxn = max(maxn, int(num))
    return maxn


def _anzahl_bilder(doc) -> int:
    xml = doc.element.xml
    inline = len(doc.inline_shapes)
    floating = xml.count("<wp:anchor")
    return inline + floating


def _hat_textbox(doc) -> bool:
    xml = doc.element.xml
    return ("txbxContent" in xml) or ("v:textbox" in xml)


# ---------------------------------------------------------------------------
# Kern-API
# ---------------------------------------------------------------------------


def pruefe(doc) -> dict:
    """Ein python-docx Document auf ATS-Risiken pruefen. {findings, stats}."""
    findings = []
    absaetze = _alle_absaetze(doc)
    body_text = "\n".join(absaetze)
    n_text = len(body_text.strip())

    # --- 1) Tabellen (FEHLER) -------------------------------------------------
    n_tabellen = len(doc.tables)
    if n_tabellen:
        findings.append(_finding(
            "tabelle", "fehler",
            str(n_tabellen) + " Tabelle(n) gefunden — ATS scrambelt/merged Zellen "
            "(Taleo/Workday). Single-Column nutzen; 'randlos' schuetzt NICHT.",
            fundstelle=str(n_tabellen) + " Tabelle(n)"))

    # --- 2) Mehrspaltiges Layout (FEHLER) -------------------------------------
    n_spalten = _spalten_anzahl(doc)
    if n_spalten > 1:
        findings.append(_finding(
            "layout", "fehler",
            "Mehrspaltiges Seitenlayout (" + str(n_spalten) + " Spalten) — ATS liest "
            "spaltenweise falsch verkettet. Einspaltig setzen.",
            fundstelle=str(n_spalten) + " Spalten"))

    # --- 3) Textbox (FEHLER) --------------------------------------------------
    if _hat_textbox(doc):
        findings.append(_finding(
            "textbox", "fehler",
            "Textbox/Shape erkannt — Inhalt wird von vielen ATS ignoriert. "
            "Text als normalen Absatz setzen."))

    # --- 4) Kontaktdaten nur in Kopf-/Fusszeile (FEHLER) ----------------------
    kf_text = _kopf_fuss_text(doc)
    kontakt_in_kf = bool(EMAIL_RE.search(kf_text) or TEL_NUM_RE.search(kf_text)
                         or TEL_RE.search(kf_text))
    kontakt_im_body = bool(EMAIL_RE.search(body_text) or TEL_NUM_RE.search(body_text)
                           or TEL_RE.search(body_text))
    if kontakt_in_kf and not kontakt_im_body:
        findings.append(_finding(
            "kontakt", "fehler",
            "Kontaktdaten stehen nur in Kopf-/Fusszeile — ATS ignoriert diese. "
            "E-Mail/Telefon als Plain-Text-Absatz oben in den Body."))
    elif not kontakt_im_body:
        findings.append(_finding(
            "kontakt", "warnung",
            "Keine Kontaktdaten (E-Mail/Telefon) als Plain-Text im Body gefunden — "
            "ATS extrahiert sie sonst nicht."))

    # --- 5) Kaum extrahierbarer Text (FEHLER) ---------------------------------
    if n_text < MIN_TEXT_LEN:
        findings.append(_finding(
            "text", "fehler",
            "Sehr wenig extrahierbarer Text (" + str(n_text) + " Zeichen) — Verdacht "
            "Scan-/Bild-CV. ATS kann nichts auslesen; als echten Text setzen.",
            fundstelle=str(n_text) + " Zeichen"))

    # --- 6) Bilder/Grafiken (WARNUNG) -----------------------------------------
    n_bilder = _anzahl_bilder(doc)
    if n_bilder:
        findings.append(_finding(
            "bild", "warnung",
            str(n_bilder) + " Bild(er)/Grafik(en) — ATS ueberspringt sie. "
            "Keine relevanten Infos nur im Bild transportieren.",
            fundstelle=str(n_bilder) + " Bild(er)"))

    # --- 7) Standard-Ueberschriften (WARNUNG) ---------------------------------
    lower = body_text.lower()
    gefundene = sorted({h for h in STANDARD_HEADINGS if h in lower})
    if not gefundene:
        findings.append(_finding(
            "ueberschrift", "warnung",
            "Keine erkennbaren Standard-Abschnittsueberschriften (Berufserfahrung/"
            "Ausbildung/Kenntnisse …) — ATS ordnet Abschnitte schlechter zu."))

    n_fehler = sum(1 for f in findings if f["schwere"] == "fehler")
    n_warn = sum(1 for f in findings if f["schwere"] == "warnung")
    return {
        "findings": findings,
        "stats": {
            "fehler": n_fehler, "warnungen": n_warn,
            "tabellen": n_tabellen, "spalten": n_spalten, "bilder": n_bilder,
            "textlaenge": n_text, "headings": len(gefundene),
            "kontakt_im_body": kontakt_im_body,
        },
    }


def pruefe_datei(pfad: str) -> dict:
    if Document is None:
        raise RuntimeError("python-docx nicht installiert (py -3.11 -m pip install python-docx)")
    return pruefe(Document(pfad))


def hat_fehler(result: dict) -> bool:
    """True, wenn mind. ein FEHLER vorliegt (Pipeline-Gate)."""
    return result["stats"]["fehler"] > 0


# ---------------------------------------------------------------------------
# CLI / Report
# ---------------------------------------------------------------------------

_SCHWERE_LABEL = {"fehler": "FEHLER", "warnung": "WARNUNG"}


def report(pfad: str, result: dict) -> str:
    s = result["stats"]
    out = []
    out.append("=== ATS-Linter — Lebenslauf-.docx auf ATS-Risiken ===")
    out.append("Datei: " + pfad)
    out.append("Struktur: " + str(s["tabellen"]) + " Tabelle(n), " + str(s["spalten"])
               + " Spalte(n), " + str(s["bilder"]) + " Bild(er), "
               + str(s["textlaenge"]) + " Zeichen Text, " + str(s["headings"])
               + " Standard-Ueberschrift(en)")
    out.append("")
    for schwere in ("fehler", "warnung"):
        gruppe = [f for f in result["findings"] if f["schwere"] == schwere]
        if not gruppe:
            continue
        out.append(_SCHWERE_LABEL[schwere] + " (" + str(len(gruppe)) + "):")
        for f in gruppe:
            out.append("  [" + f["kategorie"] + "] " + f["nachricht"])
            if f.get("fundstelle"):
                out.append("        -> " + f["fundstelle"])
        out.append("")
    if s["fehler"] == 0:
        out.append("GRUEN: keine harten ATS-Risiken (" + str(s["warnungen"])
                   + " Warnung(en) — zur Durchsicht). Ersetzt keinen echten Parser-Test.")
    else:
        out.append("ROT: " + str(s["fehler"]) + " ATS-Risiko(s), " + str(s["warnungen"])
                   + " Warnung(en) — fuer den ATS-Pfad (Single-Column) ueberarbeiten.")
    return "\n".join(out)


def main(argv):
    args = [a for a in argv[1:] if not a.startswith("--")]
    flags = {a for a in argv[1:] if a.startswith("--")}
    if not args:
        print("Aufruf: py -3.11 ats_lint.py <cv.docx> [--json]", file=sys.stderr)
        return 2
    pfad = args[0]
    if Document is None:
        print("FEHLER: python-docx nicht installiert (py -3.11 -m pip install python-docx)",
              file=sys.stderr)
        return 2
    try:
        result = pruefe_datei(pfad)
    except OSError as e:
        print("FEHLER: Datei nicht lesbar: " + str(e), file=sys.stderr)
        return 2

    if "--json" in flags:
        print(json.dumps({"datei": pfad, **result}, ensure_ascii=False, indent=2))
    else:
        print(report(pfad, result))
    return 1 if hat_fehler(result) else 0


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    sys.exit(main(sys.argv))
