# -*- coding: utf-8 -*-
"""Offline-Selbsttest fuer den QuellordnerExtractor (extract_quellordner.py).

Keine Netz-/API-Abhaengigkeit und KEIN Zugriff auf den echten PII-Quellordner:
es werden synthetische Fixtures (.odt/.docx/.pdf/.doc/Bild) in einem TEMP-Ordner
erzeugt und wieder geloescht. Prueft:
- .odt / .docx / .pdf werden korrekt extrahiert (inkl. docx-Tabellenzellen),
- zu kurze (Scan-/leer) Datei -> Status SKIP, nicht OK,
- korrupte Datei -> Status ERR, ohne den Lauf zu killen,
- .doc-Zweig robust (antiword fehlt ODER scheitert an Nicht-.doc -> SKIP, kein Crash),
- Bild/PSD wird ignoriert (nicht in der Ergebnisliste),
- corpus_text enthaelt nur OK-Dateien + Pfad-Header,
- READ-ONLY: Quell-Dateien nach Extraktion unveraendert (Hash/Size/mtime),
- READ-ONLY-Guard: Corpus-Output im Quellordner -> ValueError,
- schreibe_corpus schreibt ausserhalb + liefert plausible Statistik,
- fehlender Ordner -> NotADirectoryError.

Lauf:  py -3.11 verify_extract_quellordner.py    (exit 0 = alle Checks gruen)
"""
from __future__ import annotations

import os
import sys
import shutil
import hashlib
import zipfile
import tempfile

import extract_quellordner as ex

_checks = 0
_fails = 0


def ok(bedingung, label):
    global _checks, _fails
    _checks += 1
    if bedingung:
        print("  [OK]   " + label)
    else:
        _fails += 1
        print("  [FAIL] " + label)


# --- Fixture-Bauer ---------------------------------------------------------

_ODT_CONTENT = (
    '<?xml version="1.0" encoding="UTF-8"?>'
    '<office:document-content xmlns:office="urn:o" xmlns:text="urn:t">'
    "<office:body><office:text>"
    "<text:p>Industriemechaniker mit Fachrichtung Maschinen- und Systemtechnik bei der "
    "Karl Mayer GmbH in Obertshausen am Main.</text:p>"
    "<text:p>Taetigkeit: Guetepruefung, hydrostatische Festigkeitspruefung und Leckagepruefung "
    "von Bauteilen nach jeweiliger Vorschrift seit dem Jahr 2017, inklusive Dokumentation.</text:p>"
    "</office:text></office:body></office:document-content>"
)


def _mk_odt(pfad):
    with zipfile.ZipFile(pfad, "w") as z:
        z.writestr("content.xml", _ODT_CONTENT)


def _mk_docx(pfad, mit_tabelle=True, kurz=False):
    import docx
    d = docx.Document()
    if kurz:
        d.add_paragraph("Nur drei Worte.")
    else:
        d.add_paragraph("Bewerbung als Gueteprueferin bei der Coperion AG in Frankfurt.")
        d.add_paragraph("Erfahrung in Wareneingang, Qualitaetskontrolle und Montage nach Zeichnung.")
        if mit_tabelle:
            t = d.add_table(rows=1, cols=2)
            t.rows[0].cells[0].text = "Zeitraum: 2008 bis 2010"
            t.rows[0].cells[1].text = "Position: Pruefung am Pruefstand mit Messmitteln"
    d.save(pfad)


def _mk_pdf(pfad):
    import fitz  # pymupdf (lokal vorhanden)
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72),
                     "Arbeitszeugnis Siemens Dematic AG\n"
                     "Herr Adam hat die Foerdertechnik-Montage stets zu unserer vollsten\n"
                     "Zufriedenheit ausgefuehrt im Zeitraum von 1997 bis 2005.")
    doc.save(pfad)
    doc.close()


def _hash_dir(basis):
    """(relpath -> (size, mtime, sha1)) ueber alle Dateien — zum READ-ONLY-Nachweis."""
    snap = {}
    for wurzel, _d, dateien in os.walk(basis):
        for n in dateien:
            p = os.path.join(wurzel, n)
            with open(p, "rb") as fh:
                h = hashlib.sha1(fh.read()).hexdigest()
            st = os.stat(p)
            snap[os.path.relpath(p, basis)] = (st.st_size, int(st.st_mtime), h)
    return snap


def main():
    tmp = tempfile.mkdtemp(prefix="quellordner_verify_")
    quelle = os.path.join(tmp, "quelle")
    os.makedirs(os.path.join(quelle, "unterordner"))
    try:
        _mk_odt(os.path.join(quelle, "lebenslauf.odt"))
        _mk_docx(os.path.join(quelle, "unterordner", "anschreiben.docx"))
        _mk_pdf(os.path.join(quelle, "zeugnis.pdf"))
        _mk_docx(os.path.join(quelle, "kurz.docx"), kurz=True)        # -> SKIP
        with open(os.path.join(quelle, "kaputt.docx"), "wb") as fh:   # -> ERR
            fh.write(b"keine gueltige docx/zip-datei")
        with open(os.path.join(quelle, "fake.doc"), "wb") as fh:      # -> SKIP (antiword nonzero/fehlt)
            fh.write(b"das ist kein echtes word-doc")
        with open(os.path.join(quelle, "foto.jpg"), "wb") as fh:      # -> ignoriert
            fh.write(b"\xff\xd8\xff\xe0 jpeg")

        snap_vorher = _hash_dir(quelle)

        eintraege = ex.extrahiere(quelle)
        nach_rel = {rel: (text, status) for rel, text, status in eintraege}

        print("[1] Format-Extraktion .odt/.docx/.pdf")
        ok("lebenslauf.odt" in nach_rel and nach_rel["lebenslauf.odt"][1] == "OK", ".odt -> OK")
        ok("Karl Mayer" in nach_rel["lebenslauf.odt"][0], ".odt Inhalt extrahiert (Karl Mayer)")
        rel_docx = os.path.join("unterordner", "anschreiben.docx")
        ok(nach_rel.get(rel_docx, ("", ""))[1] == "OK", ".docx (Unterordner) -> OK")
        ok("Coperion" in nach_rel[rel_docx][0], ".docx Absatz-Inhalt extrahiert")
        ok("Pruefstand" in nach_rel[rel_docx][0], ".docx TABELLENZELLE extrahiert")
        ok(nach_rel.get("zeugnis.pdf", ("", ""))[1] == "OK", ".pdf -> OK")
        ok("Siemens" in nach_rel["zeugnis.pdf"][0], ".pdf Text-Layer extrahiert (Siemens)")

        print("[2] SKIP / ERR / ignoriert")
        ok(nach_rel["kurz.docx"][1].startswith("SKIP"), "zu kurze .docx -> SKIP")
        ok(nach_rel["kaputt.docx"][1].startswith("ERR"), "korrupte .docx -> ERR (kein Crash)")
        ok("fake.doc" in nach_rel and nach_rel["fake.doc"][1].startswith("SKIP"),
           ".doc-Zweig robust (antiword fehlt/scheitert) -> SKIP")
        ok("foto.jpg" not in nach_rel, "Bild (.jpg) wird ignoriert")

        print("[3] Konsolidierung corpus_text")
        corpus = ex.corpus_text(quelle)
        ok("### lebenslauf.odt" in corpus, "Corpus hat Pfad-Header je Datei")
        ok("Karl Mayer" in corpus and "Coperion" in corpus and "Siemens" in corpus,
           "Corpus enthaelt alle OK-Inhalte")
        ok("Nur drei Worte" not in corpus, "Corpus enthaelt KEINE SKIP-Datei")

        print("[4] READ-ONLY am Quellordner (Hash/Size/mtime unveraendert)")
        snap_nachher = _hash_dir(quelle)
        ok(snap_vorher == snap_nachher, "kein Quell-File veraendert/hinzugefuegt/geloescht")

        print("[5] READ-ONLY-Guard fuer Corpus-Output")
        guard_ausgeloest = False
        try:
            ex.schreibe_corpus(quelle, os.path.join(quelle, "corpus.txt"))
        except ValueError:
            guard_ausgeloest = True
        ok(guard_ausgeloest, "Output IM Quellordner -> ValueError (verweigert)")

        print("[6] schreibe_corpus ausserhalb + Statistik")
        out_pfad = os.path.join(tmp, "corpus_out.txt")
        st = ex.schreibe_corpus(quelle, out_pfad)
        ok(os.path.exists(out_pfad), "Corpus-Datei ausserhalb geschrieben")
        ok(st["ok"] == 3, "Statistik: 3 OK-Dateien")
        ok(st["skip"] >= 2 and st["err"] == 1, "Statistik: SKIP>=2, ERR=1")
        ok(st["corpus_zeichen"] > 0, "Statistik: corpus_zeichen > 0")

        print("[7] Fehler-Eingaben")
        fehlt = False
        try:
            ex.extrahiere(os.path.join(tmp, "gibtsnicht"))
        except NotADirectoryError:
            fehlt = True
        ok(fehlt, "fehlender Ordner -> NotADirectoryError")

    finally:
        shutil.rmtree(tmp, ignore_errors=True)

    print("\n%d Checks, %d Fehler" % (_checks, _fails))
    return 1 if _fails else 0


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    sys.exit(main())
