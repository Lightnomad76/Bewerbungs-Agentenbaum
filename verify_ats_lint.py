# -*- coding: utf-8 -*-
"""verify_ats_lint.py — offline Mechanik-Selbsttest fuer den ATS-Linter.

Baut synthetische .docx-Fixtures in-memory mit python-docx (keine PII, kein Netz)
und prueft, dass pruefe() jeden ATS-Risiko-Zweig korrekt erkennt bzw. einen sauberen
Single-Column-CV NICHT durchfallen laesst. exit 0 = alles gruen.

Lauf: py -3.11 verify_ats_lint.py   (braucht python-docx -> global py -3.11, NICHT venv)
"""
from __future__ import annotations

import sys
import zlib
import struct
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ats_lint import pruefe, hat_fehler

_fehler = 0


def _png_1x1() -> bytes:
    """Garantiert gueltiges 1x1-RGB-PNG (programmatisch, kein fragiles Base64-Literal)."""
    def chunk(typ, data):
        roh = typ + data
        return struct.pack(">I", len(data)) + roh + struct.pack(">I", zlib.crc32(roh) & 0xffffffff)
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))  # 1x1, 8-bit, RGB
    idat = chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))            # Filter 0 + 1 Pixel
    iend = chunk(b"IEND", b"")
    return sig + ihdr + idat + iend


_PNG = _png_1x1()


def ok(bedingung, text, ist=None):
    global _fehler
    marke = "[OK  ]" if bedingung else "[FAIL]"
    extra = "" if bedingung else ("  (ist: " + repr(ist) + ")")
    print("  " + marke + "   " + text + extra)
    if not bedingung:
        _fehler += 1


def _hat(result, kategorie, schwere=None):
    return any(f["kategorie"] == kategorie and (schwere is None or f["schwere"] == schwere)
               for f in result["findings"])


def basis_cv():
    """Sauberer Single-Column-CV: Kontakt im Body, Standard-Ueberschriften, genug Text."""
    d = Document()
    d.add_paragraph("Max Mustermann")
    d.add_paragraph("max.mustermann@example.com | 0151 23456789")
    d.add_paragraph("Berufserfahrung")
    d.add_paragraph("2018-2024 Industriemechaniker bei der Beispiel GmbH. Wartung und "
                    "Instandhaltung von CNC-Anlagen, Montage von Baugruppen, Fehlersuche "
                    "an Hydraulik- und Pneumatiksystemen sowie Qualitaetssicherung.")
    d.add_paragraph("Ausbildung")
    d.add_paragraph("2014-2018 Ausbildung zum Industriemechaniker, Beispielstadt.")
    d.add_paragraph("Kenntnisse")
    d.add_paragraph("CNC, Drehen, Fraesen, SPS-Grundlagen, MS Office, technisches Zeichnen.")
    return d


def setze_spalten(d, n):
    sectPr = d.sections[0]._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(n))


def main():
    print("[1] Sauberer Single-Column-CV faellt NICHT durch")
    res = pruefe(basis_cv())
    ok(res["stats"]["fehler"] == 0, "0 FEHLER", res["stats"])
    ok(res["stats"]["warnungen"] == 0, "0 WARNUNG", res["stats"])
    ok(hat_fehler(res) is False, "hat_fehler False")
    ok(res["stats"]["kontakt_im_body"] is True, "Kontakt im Body erkannt")
    ok(res["stats"]["headings"] >= 3, "Standard-Ueberschriften gezaehlt", res["stats"]["headings"])

    print("[2] Tabelle = FEHLER")
    d = basis_cv()
    d.add_table(rows=2, cols=2)
    r = pruefe(d)
    ok(_hat(r, "tabelle", "fehler"), "Tabelle als FEHLER")
    ok(r["stats"]["tabellen"] == 1, "Tabellen-Zaehler = 1", r["stats"]["tabellen"])
    ok(hat_fehler(r) is True, "Gate rot")

    print("[3] Mehrspaltiges Layout = FEHLER")
    d = basis_cv()
    setze_spalten(d, 2)
    r = pruefe(d)
    ok(_hat(r, "layout", "fehler"), "Mehrspaltig als FEHLER")
    ok(r["stats"]["spalten"] == 2, "Spalten-Zaehler = 2", r["stats"]["spalten"])

    print("[4] Kontakt nur in Kopfzeile = FEHLER")
    d = Document()
    d.add_paragraph("Berufserfahrung")
    d.add_paragraph("2018-2024 Industriemechaniker bei der Beispiel GmbH mit Wartung, "
                    "Instandhaltung, Montage und Qualitaetssicherung an CNC-Anlagen.")
    d.add_paragraph("Ausbildung")
    d.add_paragraph("2014-2018 Ausbildung zum Industriemechaniker in Beispielstadt.")
    d.sections[0].header.paragraphs[0].text = "max.mustermann@example.com | 0151 23456789"
    r = pruefe(d)
    ok(_hat(r, "kontakt", "fehler"), "Kontakt-nur-im-Header als FEHLER")
    ok(r["stats"]["kontakt_im_body"] is False, "kein Kontakt im Body")

    print("[5] Kaum Text (Scan/Bild-CV) = FEHLER")
    d = Document()
    d.add_paragraph("Lebenslauf")
    r = pruefe(d)
    ok(_hat(r, "text", "fehler"), "Zu wenig Text als FEHLER")
    ok(r["stats"]["textlaenge"] < 200, "Textlaenge unter Schwelle", r["stats"]["textlaenge"])

    print("[6] Eingebettetes Bild = WARNUNG")
    d = basis_cv()
    d.add_picture(BytesIO(_PNG))
    r = pruefe(d)
    ok(_hat(r, "bild", "warnung"), "Bild als WARNUNG")
    ok(r["stats"]["bilder"] >= 1, "Bild-Zaehler >= 1", r["stats"]["bilder"])
    ok(hat_fehler(r) is False, "Bild allein kein FEHLER")

    print("[7] Keine Standard-Ueberschriften = WARNUNG")
    d = Document()
    d.add_paragraph("kontakt@example.com 0151 23456789")
    d.add_paragraph("Eine ausreichend lange Fliesstext-Beschreibung ohne jegliche "
                    "typische Rubrik-Bezeichnung, damit ausschliesslich die fehlende "
                    "Gliederung bemaengelt wird und nicht die Textlaenge unterschritten.")
    r = pruefe(d)
    ok(_hat(r, "ueberschrift", "warnung"), "fehlende Ueberschriften als WARNUNG")
    ok(r["stats"]["fehler"] == 0, "kein FEHLER in diesem Fall", r["stats"])

    print("[8] Headings in Tabellen-Zellen werden mitgelesen")
    d = Document()
    d.add_paragraph("kontakt@example.com 0151 23456789")
    t = d.add_table(rows=1, cols=1)
    t.rows[0].cells[0].paragraphs[0].text = ("Berufserfahrung: 2018-2024 Industriemechaniker "
                                             "mit Wartung und Instandhaltung an CNC-Anlagen.")
    r = pruefe(d)
    ok(r["stats"]["headings"] >= 1, "Heading aus Tabellen-Zelle erkannt", r["stats"]["headings"])
    ok(_hat(r, "tabelle", "fehler"), "Tabelle trotzdem als FEHLER")

    print("[9] Schema-Stabilitaet")
    res = pruefe(basis_cv())
    ok(set(res.keys()) == {"findings", "stats"}, "Top-Level-Keys stabil", set(res.keys()))
    ok(set(res["stats"].keys()) == {"fehler", "warnungen", "tabellen", "spalten", "bilder",
                                    "textlaenge", "headings", "kontakt_im_body"},
       "stats-Keys stabil", set(res["stats"].keys()))
    d = basis_cv()
    d.add_table(rows=1, cols=1)
    f = pruefe(d)["findings"][0]
    ok(set(f.keys()) == {"kategorie", "schwere", "nachricht", "fundstelle"},
       "finding-Schema stabil", set(f.keys()))

    print("")
    if _fehler == 0:
        print("GRUEN: alle ATS-Linter-Checks bestanden")
        return 0
    print("ROT: " + str(_fehler) + " Check(s) fehlgeschlagen")
    return 1


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    sys.exit(main())
